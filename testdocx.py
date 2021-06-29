# %%
from docx import Document
from docx.shared import Inches
from docx.table import Table

# %% 
docu = Document("C:/Program Files/pdfbox/2020Q1.docx")
docu


# %%
len(docu.tables)
# %%
type(docu.tables[10])

# %%
table = docu.tables[7]
table = Table()


# %%
table.cell(0,0).text
# %%
len(docu.paragraphs)

# %%

# %%
for parag in docu.paragraphs:
    print(parag.text)
# %%
print(docu.paragraphs[149].text)
# %%
